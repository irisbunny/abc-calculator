import streamlit as st
import random
import math
from datetime import datetime
from docx import Document
from docx.shared import Pt
from io import BytesIO
import datetime as dt

# 设置页面
st.set_page_config(page_title="智能计算系统", layout="centered", initial_sidebar_state="expanded")

# 自定义样式
st.markdown("""
<style>
    .stNumberInput input, .stTextInput input, .stDateInput input {
        border-radius: 8px !important;
        border: 1px solid #ced4da !important;
    }
    .stButton>button {
        background-color: #4a6bdf !important;
        color: white !important;
        transition: all 0.3s !important;
    }
    .stButton>button:hover {
        background-color: #3a5bc7 !important;
        transform: scale(1.02);
    }
    .result-box {
        background: white;
        border-radius: 10px;
        padding: 15px;
        margin: 10px 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    [data-testid="stSidebar"] {
        background: #f8f9fa !important;
    }
</style>
""", unsafe_allow_html=True)

# =============== ABC计算器功能 ===============
def abc_calculator():
    st.title("🧮 ABC组合计算器")
    
    # 优化后的计算函数
    def find_abc_combinations(d, max_results=5, max_attempts=10000000):
        results = []
        attempt = 0
        
        while len(results) < max_results and attempt < max_attempts:
            b = round(random.uniform(2000, 5000), 1)
            a = round(random.uniform(1, 5), 1)
            c = round(random.uniform(0.5, 5), 1)
            
            d_candidate = a * b * c
            
            if math.floor(d_candidate * 10) / 10 == d:
                results.append({
                    'a': a,
                    'b': b,
                    'c': c,
                    'd_val': d_candidate,
                    'floor': math.floor(d_candidate * 10) / 10
                })
            attempt += 1
        
        return results, attempt

    # 侧边栏设置
    with st.sidebar:
        st.header("⚙️ 计算设置")
        max_results = st.slider("需要的结果数量", 1, 10, 5, key='max_results')
        max_attempts = st.number_input("最大尝试次数", 1000, 10000000, 1000000, step=10000)

    # 主界面
    d_value = st.number_input("目标d值", min_value=0.1, value=None, step=0.1, format="%.1f")

    if st.button("🚀 开始计算", key="abc_calculate"):
        if d_value is None:
            st.warning("请输入目标d值")
        else:
            with st.spinner(f'正在搜索 {max_results} 组解...'):
                start_time = datetime.now()
                results, total_attempts = find_abc_combinations(d_value, max_results, max_attempts)
                compute_time = (datetime.now() - start_time).total_seconds()
                
                if not results:
                    st.error(f"未找到解 (尝试了 {total_attempts} 次)")
                else:
                    st.success(f"找到 {len(results)} 组解 | 耗时 {compute_time:.2f}s)")
                    
                    for i, sol in enumerate(results, 1):
                        with st.expander(f"组合 {i}", expanded=True):
                            st.markdown(f"""
                            <div class="result-box">
                                <p>🔢 <b>参数</b>: a={sol['a']:.1f}, b={sol['b']:.1f}, c={sol['c']:.1f}</p>
                                <p>🧮 <b>计算值</b>: {sol['d_val']:.4f}</p>
                                <p>⌛ <b>向下取整</b>: {sol['floor']}</p>
                            </div>
                            """, unsafe_allow_html=True)

# =============== 结算单生成器功能 ===============
def invoice_generator():
    st.title("📄 结算单生成系统")
    
    # 结算单模板生成函数
    from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from io import BytesIO

def generate_invoice(party_a, party_b, contract_date, start_date, end_date, a, b, c, d):
    doc = Document()

    # 设置默认字体：宋体 10.5pt
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置页面边距
    section = doc.sections[0]
    section.top_margin = Inches(2.61 / 2.54)
    section.bottom_margin = Inches(2.08 / 2.54)
    section.left_margin = Inches(3.00 / 2.54)
    section.right_margin = Inches(2.75 / 2.54)

    # 添加标题
    title = doc.add_paragraph("附件二 结算单", style='Heading 1')
    title_format = title.paragraph_format
    title_format.space_before = Pt(2)
    title_format.space_after = Pt(0)

    # 添加表格（6行5列）
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'

    # 设置列宽（单位EMU）
    column_widths = [1617345, 810260, 899160, 899160, 989330]
    for row in table.rows:
        for i, width in enumerate(column_widths):
            row.cells[i].width = width

    # 单元格填充函数
    def fill_cells(row, col, text, col_span=1):
        cell = table.cell(row, col)
        cell.text = text

        para = cell.paragraphs[0]
        run = para.runs[0]
        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(0)

        # 合并单元格（横向）
        if col_span > 1:
            cell.merge(table.cell(row, col + col_span - 1))

    # 填充内容
    fill_cells(0, 0, "甲方")
    fill_cells(0, 1, party_a, col_span=4)

    fill_cells(1, 0, "乙方")
    fill_cells(1, 1, party_b, col_span=4)

    fill_cells(2, 0, "合作内容")
    fill_cells(2, 1, f"根据甲乙双方于{contract_date}签署的《委托开发及运维服务外包协议》（简称：主协议），甲方为乙方提供如下技术服务，乙方支付费用。", col_span=4)

    fill_cells(3, 0, "结算周期")
    fill_cells(3, 1, f"{start_date}至{end_date}", col_span=4)

    fill_cells(4, 0, "技术支持服务")
    fill_cells(4, 1, "服务人次")
    fill_cells(4, 2, "天单价")
    fill_cells(4, 3, "人天数")
    fill_cells(4, 4, "结算金额")

    fill_cells(5, 0, "人工成本")
    fill_cells(5, 1, str(a))
    fill_cells(5, 2, str(b))
    fill_cells(5, 3, str(c))
    fill_cells(5, 4, str(d))

    # 导出为内存文件
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


    # 表单设计
    with st.form("invoice_form"):
        col1, col2 = st.columns(2)
        with col1:
            party_a = st.text_input("甲方名称", "钜韜有限公司")
            party_b = st.text_input("乙方名称", "xx公司")
            contract_date = st.date_input("协议签署日期", dt.date(2025, 1, 1))
        
        with col2:
            start_date = st.date_input("结算开始日期", dt.date(2025, 3, 1))
            end_date = st.date_input("结算结束日期", dt.date(2025, 3, 31))
            
            col_a, col_b = st.columns(2)
            with col_a:
                a = st.number_input("服务人次(a)", min_value=1, value=1)
                b = st.number_input("天单价(b)", min_value=0, value=1000)
            with col_b:
                c = st.number_input("人天数(c)", min_value=1, value=30)
                d = st.number_input("结算金额(d)", min_value=0, value=30000)
        
        submitted = st.form_submit_button("生成结算单")

    if submitted:
        # 生成Word文档
        file_stream = generate_invoice(
            party_a=party_a,
            party_b=party_b,
            contract_date=contract_date.strftime("%Y/%m/%d"),
            start_date=start_date.strftime("%Y/%m/%d"),
            end_date=end_date.strftime("%Y/%m/%d"),
            a=a, b=b, c=c, d=d
        )
        
        # 提供下载
        st.success("结算单生成成功！")
        st.download_button(
            label="⬇️ 下载结算单",
            data=file_stream,
            file_name=f"结算单_{dt.datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# =============== 主程序 ===============
# 侧边栏导航
st.sidebar.title('功能导航')
app_mode = st.sidebar.radio("", 
    ["ABC组合计算器", "结算单生成器"],
    index=0  # 默认选中第一个
)

# 根据选择显示不同页面
if app_mode == "ABC组合计算器":
    abc_calculator()
else:
    invoice_generator()

# 侧边栏底部信息
st.sidebar.markdown("---")
st.sidebar.info(
    "💡 使用提示：\n"
    "- ABC计算器：输入d值查找组合\n"
    "- 结算单生成器：填写表单自动生成文档"
)
