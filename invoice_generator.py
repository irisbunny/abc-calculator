import streamlit as st
from docx import Document
from docx.shared import Pt
from io import BytesIO
import datetime

# 设置页面
st.set_page_config(page_title="结算单生成器", layout="centered")

# 自定义样式
st.markdown("""
<style>
    .stTextInput input, .stDateInput input, .stNumberInput input {
        border-radius: 8px !important;
        border: 1px solid #ced4da !important;
    }
    .stButton>button {
        background-color: #4a6bdf !important;
        color: white !important;
    }
</style>
""", unsafe_allow_html=True)

# 结算单模板生成函数
def generate_invoice(party_a, party_b, contract_date, start_date, end_date, a, b, c, d):
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    
    # 添加标题
    doc.add_paragraph('附件二 结算单', style='Heading 1')
    
    # 创建表格
    table = doc.add_table(rows=7, cols=5)
    table.style = 'Table Grid'
    
    # 填充表格内容
    cells = table.rows[0].cells
    cells[0].text = "甲方"
    cells[1].text = party_a
    cells[1].merge(cells[4])
    
    # 乙方信息
    cells = table.rows[1].cells
    cells[0].text = "乙方"
    cells[1].text = party_b
    cells[1].merge(cells[4])
    
    # 合作内容
    cells = table.rows[2].cells
    cells[0].text = "合作内容"
    cells[1].text = f"根据甲乙双方于{contract_date}签署的《委托开发及运维服务外包协议》（简称：主协议），甲方为乙方提供如下技术服务，乙方支付费用。"
    cells[1].merge(cells[4])
    
    # 结算周期
    cells = table.rows[3].cells
    cells[0].text = "结算周期"
    cells[1].text = f"{start_date}至{end_date}"
    cells[1].merge(cells[4])
    
    # 表头
    cells = table.rows[4].cells
    cells[0].text = "技术支持服务"
    cells[1].text = "服务人次"
    cells[2].text = "天单价"
    cells[3].text = "人天数"
    cells[4].text = "结算金额"
    
    # 数据行
    cells = table.rows[5].cells
    cells[0].text = "人工成本"
    cells[1].text = str(a)
    cells[2].text = str(b)
    cells[3].text = str(c)
    cells[4].text = str(d)
    
    # 保存到内存
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream

# 主界面
st.title("📄 结算单生成系统")

with st.form("invoice_form"):
    col1, col2 = st.columns(2)
    with col1:
        party_a = st.text_input("甲方名称", "钜韜有限公司")
        party_b = st.text_input("乙方名称", "xx公司")
        contract_date = st.date_input("协议签署日期", datetime.date(2025, 1, 1))
    
    with col2:
        start_date = st.date_input("结算开始日期", datetime.date(2025, 3, 1))
        end_date = st.date_input("结算结束日期", datetime.date(2025, 3, 31))
        
        col_a, col_b = st.columns(2)
        with col_a:
            a = st.number_input("服务人次(a)", min_value=1, value=1)
            b = st.number_input("天单价(b)", min_value=0, value=1000)
        with col_b:
            c = st.number_input("人天数(c)", min_value=1, value=30)
            d = st.number_input("结算金额(d)", min_value=0, value=30000)
    
    submitted = st.form_submit_button("生成结算单")

if submitted:
    # 生成Word文档
    file_stream = generate_invoice(
        party_a=party_a,
        party_b=party_b,
        contract_date=contract_date.strftime("%Y/%m/%d"),
        start_date=start_date.strftime("%Y/%m/%d"),
        end_date=end_date.strftime("%Y/%m/%d"),
        a=a, b=b, c=c, d=d
    )
    
    # 提供下载
    st.success("结算单生成成功！")
    st.download_button(
        label="⬇️ 下载结算单",
        data=file_stream,
        file_name=f"结算单_{datetime.datetime.now().strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    # 预览区域
    st.subheader("预览效果")
    st.image("https://via.placeholder.com/600x400?text=结算单预览图", caption="结算单预览")
